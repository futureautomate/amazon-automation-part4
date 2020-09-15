# -*- coding: utf-8 -*-
"""
Created on Sun Mar 29 23:51:04 2020

@author: Tejas
"""

from selenium import webdriver
import time
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
import logindata
import urllib.request
import docx

options = webdriver.ChromeOptions()
options.add_argument("--start-maximized")
driver = webdriver.Chrome('F:\Channel\webdriver\chromedriver.exe', chrome_options=options)
action = ActionChains(driver)
time.sleep(1)


driver.get('http://www.amazon.in')
time.sleep(3)
 
firstLevelMenu = driver.find_element_by_xpath('//*[@id="nav-link-accountList"]/span/span')
action.move_to_element(firstLevelMenu).perform()
time.sleep(3)
 
secondLevelMenu = driver.find_element_by_xpath('//*[@id="nav-flyout-ya-signin"]/a/span')
secondLevelMenu.click()
time.sleep(3)

signinelement = driver.find_element_by_xpath('//*[@id="ap_email"]')
signinelement.send_keys(logindata.USERNAME)
time.sleep(3)

cont = driver.find_element_by_xpath('//*[@id="continue"]')
cont.click()
time.sleep(3)

passwordelement = driver.find_element_by_xpath('//*[@id="ap_password"]')
passwordelement.send_keys(logindata.PASSWORD)
time.sleep(3)

login = driver.find_element_by_xpath('//*[@id="signInSubmit"]')
login.click()
time.sleep(3)

'-----------------------------XXXX-----------------------------------------'

textBox = driver.find_element_by_id('twotabsearchtextbox')
textBox.send_keys('iphone 6s')
textBox.send_keys(Keys.ENTER)

mydoc = docx.Document()
i = 1

productPage = driver.find_element_by_css_selector('div.s-main-slot')
productLists = productPage.find_elements_by_css_selector('div.s-latency-cf-section')

for product in productLists:  
    
    productDetails = product.find_element_by_css_selector('h2.a-size-mini')
    
    try:       
        rating = product.find_element_by_css_selector('div.a-size-small')
        rate = rating.find_element_by_tag_name('span')
        
    except:
        continue
       
    image = product.find_element_by_tag_name('img')
    urllib.request.urlretrieve(image.get_attribute('src'),"images/{}.jpg".format(i))
    mydoc.add_heading(productDetails.text, 0)
    mydoc.add_picture("images/{}.jpg".format(i), width=docx.shared.Inches(2), height=docx.shared.Inches(4))
    mydoc.add_heading(rate.get_attribute('aria-label'), 1)
    
    mydoc.add_page_break()
    
    i+=1
    
mydoc.save("amazon_data.docx")
        


